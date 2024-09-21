import streamlit as st
import requests
import os
from dotenv import load_dotenv
import json
import random
from PyPDF2 import PdfReader
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Inches
import base64
import pytesseract
from PIL import Image

# Load environment variables
load_dotenv()

# API keys
SERPER_API_KEY = os.getenv("SERPER_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
QUIZAPI_KEY = os.getenv("QUIZAPI_KEY")

def search_serper(query):
    url = "https://google.serper.dev/search"
    
    payload = json.dumps({"q": query})
    headers = {
        'X-API-KEY': SERPER_API_KEY,
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return response.json()

def query_groq(prompt):
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "llama-3.1-70b-versatile",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.5,
        "max_tokens": 1000
    }
    response = requests.post(url, headers=headers, json=data)
    return response.json()['choices'][0]['message']['content']

def get_quiz_questions():
    url = "https://quizapi.io/api/v1/questions"
    params = {
        "apiKey": QUIZAPI_KEY,
        "limit": 5,
        "category": "law",
        "difficulty": "Medium"
    }
    response = requests.get(url, params=params)
    return response.json()

def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

def main():
    local_css("styles.css")
    
    st.markdown('<style>body {font-family: "Graduate", sans-serif !important;}</style>', unsafe_allow_html=True)
    
    st.markdown("<h1 class='main-title'>Better Chat Advocate</h1>", unsafe_allow_html=True)
    
    features = [
        "Legal Research", "Case Documents Analysis", "Legal Quiz", "Generate a Complaint",
        "Case Prediction", "Legal Process Navigator", "Document Templates",
        "Legal Aid", "Legal News", "Constitutional History of India", "Case Tracker", "Rights Awareness",
        "Cyber Law Awareness", "Environmental Law Information", "Property Law Guide", "Personal Law Guide"  # New features
    ]
    
    if 'selected_feature' not in st.session_state:
        st.session_state.selected_feature = None

    cols = st.columns(2)
    for i, feature in enumerate(features):
        if cols[i % 2].button(feature, key=f"btn_{feature}"):
            st.session_state.selected_feature = feature

    if st.session_state.selected_feature:
        handle_feature(st.session_state.selected_feature)
    
    # Sidebar
    with st.sidebar:
        st.markdown("""
        <div style="font-family: 'Graduate', sans-serif; background-color: #000080; color: white; padding: 20px; border-radius: 10px; border: 2px solid #1E90FF;">
            <h2 style="color: white; font-size: 3.0em;">Legal Assistance</h2>
        </div>
        """, unsafe_allow_html=True)
        st.markdown('<div class="sidebar-content">', unsafe_allow_html=True)
        
        st.sidebar.markdown("---")
        if st.sidebar.button("About", key="about", help="Learn more", use_container_width=True):
            st.sidebar.markdown("""
            <span style='color: red;'>**Better Chat Advocate:**</span> Law can be like a twisted game of chess - one wrong move and you're in checkmate. But don't sweat it, we've got your back. At Better Chat Advocate, we're not just about advocating for your rights, we're about empowering you with knowledge. Think of us as your personal guide through the tangled web of legalese. We break down the complexities, so you can make informed decisions. Educate yourself on the ways of the law, and learn how to draft legal complaints like a pro. With us, you're not just fighting a case, you're taking back control. The purpose of Better Chat Advocate is educate the Indian citizens with regards to their rights and the law.
            """, unsafe_allow_html=True)
        
        # Unified style for all sidebar buttons
        button_style = """
        <style>
        .sidebar .sidebar-content .stButton > button {
            width: 100% !important;
            height: 50px !important;
            background-color: #000080 !important;
            color: white !important;
            font-size: 16px !important;
            font-weight: bold !important;
            border: none !important;
            border-radius: 5px !important;
            margin-bottom: 10px !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
        }
        .sidebar .sidebar-content .stButton[data-testid="About"] > button,
        .sidebar .sidebar-content .stButton[data-testid="How to Use?"] > button,
        .sidebar .sidebar-content .stButton[data-testid="Contact Us"] > button {
            background-color: #000000 !important;
        }
        </style>
        """
        st.markdown(button_style, unsafe_allow_html=True)
        
        # Remove wrapper divs and use consistent button creation
        how_to_use_button = st.sidebar.button("How to Use?", key="how_to_use", help="Best Practices", use_container_width=True)
        contact_button = st.sidebar.button("Contact Us", key="contact_us", help="Get in touch", use_container_width=True)
        
        # Handle button clicks
        if st.session_state.get("how_to_use"):
            st.sidebar.markdown("""
            Now that you've got <span style='color: red;'><strong>Better Chat Advocate</strong></span>, let's get you started. Here's how it works:
            
            • Open the app and start chatting with me - it's that simple!
            • I'll ask you some questions, and you respond with the details of your issue.
            • My AI agents will analyze your problem and provide personalized advice.
            • You can ask questions, describe issues, or get second opinions.
            • You'll receive actionable advice, clear explanations (no lawyer speak), and next steps.
            • I'll walk you through the process, so you know what to expect.
            • No fuss, no muss - just straightforward advice from a trusted advocate.
           
            
            Get started today, and let's take on your legal issues together!
            """, unsafe_allow_html=True)
        
        if st.session_state.get("contact_us"):
            st.sidebar.markdown("""
            • <span style='color: red;'>**Better Chat Advocate**</span> is just the start, it's a guiding light, but it's not a substitute for real deal, professional advice.
            • If you truly feel like your rights are being railroaded, don't mess around, seek out a trusted lawyer or authority figure who can help you navigate the system.
            • This app is meant to educate, inform, and empower you, but it's not a replacement for human expertise.
            • If you're dealing with a serious issue, don't hesitate, get the help you need today.
            • If you've got any questions or concerns about the app itself, shoot me an email at ggengineerco@gmail.com.
            • <span style='color: red;'>From Engineer.
            """, unsafe_allow_html=True)
        
        # Added animated waving Indian flag below Contact Us, fixed at bottom left
        st.sidebar.markdown("""
            <div style="width: 100%; overflow: hidden;">
                <img src="https://media.giphy.com/media/9Gnbm29r7ftUA/giphy.gif" style="width: 100%; object-fit: cover;">
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
            <style>
            .sidebar .sidebar-content {
                background-color: #FFF5E6;
            }
            .sidebar .sidebar-content .stButton > button {
                width: 100%;
            }
            .stButton > button {
                background-color: #000080 !important;
                color: #FFFFFF !important;
            }
            .stTextInput > div > div > input {
                background-color: #000000 !important;
                color: #FFFFFF !important;
            }
            .stTextArea > div > div > textarea {
                background-color: #000000 !important;
                color: #FFFFFF !important;
            }
            .stSelectbox > div > div > select {
                background-color: #000000 !important;
                color: #FFFFFF !important;
            }
            body {
                color: #FFFFFF !important;
            }
            </style>
        """, unsafe_allow_html=True)

def handle_feature(feature):
    if feature == "Case Documents Analysis":
        document_analysis()
    elif feature == "Case Prediction":
        case_prediction()
    elif feature == "Legal Quiz":
        legal_quiz()
    elif feature == "Generate a Complaint":
        complaint_generator()
    elif feature == "Legal Research":
        legal_research()
    elif feature == "Rights Awareness":
        rights_awareness()
    elif feature == "Legal Process Navigator":
        legal_process_navigator()
    elif feature == "Document Templates":
        document_templates()
    elif feature == "Legal News":
        legal_news()
    elif feature == "Legal Aid":
        legal_aid()
    elif feature == "Case Tracker":
        case_tracker()
    elif feature == "Constitutional History of India":
        constitutional_history()
    elif feature == "Personal Law Guide":  # New handler
        personal_law_guide()
    elif feature == "Environmental Law Information":  # New handler
        environmental_law_info()
    elif feature == "Cyber Law Awareness":  # New handler
        cyber_law_awareness()
    elif feature == "Property Law Guide":  # New handler
        property_law_guide()

def document_analysis():
    st.subheader("Legal Document Analysis")
    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "png", "jpg", "jpeg"])
    if uploaded_file is not None:
        text = ""
        if uploaded_file.type == "application/pdf":
            pdf_reader = PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                text += page.extract_text()
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif uploaded_file.type.startswith("image/"):
            image = Image.open(uploaded_file)
            text = pytesseract.image_to_string(image)
        
        prompt = f"Analyze the following legal document and provide a summary, key entities, and main legal points. Only provide results for legal analysis:\n\n{text[:4000]}"
        analysis = query_groq(prompt)
        st.write(analysis)

def case_prediction():
    st.subheader("Case Prediction and Analysis")
    case_description = st.text_area("Enter case details:")
    if st.button("Predict"):
        search_results = search_serper(f"Indian law cases similar to: {case_description}")
        prompt = f"Based on the following case description and similar cases, predict the outcome and provide an analysis. Based on historical data, provide your key judgement as the best lawyer and provide the steps for the user to follow. Include crucial information and key points:\n\nCase: {case_description}\n\nSimilar cases: {json.dumps(search_results['organic'][:3])}"
        prediction = query_groq(prompt)
        st.write(prediction)

def legal_quiz():
    st.subheader("Interactive Legal Quiz")
    
    law_areas = ["Constitutional Law", "Criminal Law", "Family Law", "Property Law", "Contract Law", "Corporate Law"]
    selected_area = st.selectbox("Select an area of Indian law:", law_areas)
    
    if 'quiz_started' not in st.session_state:
        st.session_state.quiz_started = False
    
    if 'quiz_data' not in st.session_state:
        st.session_state.quiz_data = []
    
    if 'user_answers' not in st.session_state:
        st.session_state.user_answers = {}
    
    if 'quiz_ended' not in st.session_state:
        st.session_state.quiz_ended = False
    
    if st.button("Start Quiz") or (st.session_state.quiz_ended and st.button("Start New Quiz")):
        st.session_state.quiz_started = True
        st.session_state.quiz_ended = False
        st.session_state.quiz_data = []
        st.session_state.user_answers = {}
        
        search_results = search_serper(f"Indian {selected_area} key concepts")
        
        prompt = f"""Based on the following information about Indian {selected_area}, generate 5 multiple-choice questions. 
        For each question, provide the question text, 4 options (with one correct answer), and indicate the correct answer.
        Format each question as follows:
        Q: [Question text]
        A) [Option A]
        B) [Option B]
        C) [Option C]
        D) [Option D]
        Correct: [A/B/C/D]

        Information:
        {json.dumps(search_results['organic'][:3])}"""
        
        quiz_data_str = query_groq(prompt)
        
        quiz_data = []
        current_question = {}
        for line in quiz_data_str.split('\n'):
            line = line.strip()
            if line.startswith('Q:'):
                if current_question:
                    quiz_data.append(current_question)
                current_question = {'question': line[2:].strip(), 'options': []}
            elif line.startswith(('A)', 'B)', 'C)', 'D)')):
                current_question['options'].append(line[2:].strip())
            elif line.startswith('Correct:'):
                current_question['correct_answer'] = ord(line[-1]) - ord('A')
        if current_question:
            quiz_data.append(current_question)
        
        if len(quiz_data) == 0:
            st.error("Failed to generate quiz questions. Please try again.")
        else:
            st.session_state.quiz_data = quiz_data
    
    if st.session_state.quiz_started and not st.session_state.quiz_ended:
        for i, q in enumerate(st.session_state.quiz_data):
            st.write(f"Q{i+1}: {q['question']}")
            user_answer = st.radio(f"Select your answer for Q{i+1}", q['options'], key=f"q{i}")
            st.session_state.user_answers[i] = user_answer
        
        if st.button("End Test"):
            st.session_state.quiz_ended = True
            score = 0
            for i, q in enumerate(st.session_state.quiz_data):
                user_answer = st.session_state.user_answers.get(i)
                if user_answer == q['options'][q['correct_answer']]:
                    score += 1
            
            percentage = (score / len(st.session_state.quiz_data)) * 100
            st.write(f"Your score: {score}/{len(st.session_state.quiz_data)} ({percentage:.2f}%)")
            
            for i, q in enumerate(st.session_state.quiz_data):
                st.write(f"Q{i+1}: {q['question']}")
                user_answer = st.session_state.user_answers.get(i)
                correct_answer = q['options'][q['correct_answer']]
                if user_answer == correct_answer:
                    st.success(f"Your answer: {user_answer} (Correct)")
                else:
                    st.error(f"Your answer: {user_answer}")
                    st.info(f"Correct answer: {correct_answer}")
    
    if st.session_state.quiz_ended:
        st.session_state.quiz_started = False

def complaint_generator():
    st.subheader("Complaint Generator")
    name = st.text_input("Your Name")
    incident_date = st.date_input("Date of Incident")
    incident_description = st.text_area("Describe the incident")
    if st.button("Generate Complaint"):
        prompt = f"Generate a formal legal complaint for Indian Law systems, based on the following information:\nName: {name}\nDate of Incident: {incident_date}\nDescription: {incident_description}"
        complaint = query_groq(prompt)
        st.write(complaint)
        
        docx = create_docx(complaint)
        st.download_button(
            label="Download as DOCX",
            data=docx,
            file_name="complaint.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

def legal_research():
    st.subheader("Legal Research Assistant")
    query = st.text_input("Enter your legal research query")
    if st.button("Search"):
        search_results = search_serper(f"Indian law: {query}")
        prompt = f"Summarize the following legal research results for the query. Ensure that latest data is provided for the user query and all information is provided in a structured format so that the user is educated. '{query}':\n\n{json.dumps(search_results['organic'][:5])}"
        summary = query_groq(prompt)
        st.write(summary)

def rights_awareness():
    st.subheader("Rights Awareness Module")
    rights_category = st.selectbox("Select a category", ["Fundamental Rights", "Consumer Rights", "Labor Rights", "Constitutional Rights", "Women's Rights"])
    search_results = search_serper(f"Indian {rights_category} explanation")
    prompt = f"Provide a comprehensive explanation of {rights_category} in India based on the following search results. Provide the most important information in bullets and follow up with any other material that is relevant. Dictate the rights as a rights anchor and refrain from using terms like *here are the 10 bullets as asked*. Begin and end with actual relevant info.:\n\n{json.dumps(search_results['organic'][:5])}"
    explanation = query_groq(prompt)
    st.write(explanation)

def personal_law_guide():
    st.subheader("Personal Law Guide")
    law_category = st.selectbox("Select a category", ["Hindu Law", "Muslim Law", "Christian Law", "Parsi Law", "Special Marriage Act", "Personal Finance Law"])
    search_results = search_serper(f"Indian {law_category} explanation")
    prompt = f"Provide a comprehensive explanation of {law_category} in India based on the following search results. Provide the most important information in bullets and follow up with any other material that is relevant. Dictate the information as a information anchor and refrain from using terms like *here are the 10 bullets as asked*. Begin and end with actual relevant info.:\n\n{json.dumps(search_results['organic'][:5])}"
    explanation = query_groq(prompt)
    st.write(explanation)

def environmental_law_info():
    st.subheader("Environmental Law Information")
    env_topic = st.selectbox("Select a topic", ["Environmental Protection Act", "Wildlife Protection Act", "Forest Conservation Act", "Air Pollution Control", "Water Pollution Control"])
    search_results = search_serper(f"Indian {env_topic} explanation")
    prompt = f"Provide a comprehensive explanation of {env_topic} in India based on the following search results with accurate and real time data:\n\n{json.dumps(search_results['organic'][:5])}"
    explanation = query_groq(prompt)
    st.write(explanation)

def cyber_law_awareness():
    st.subheader("Cyber Law Awareness")
    cyber_topic = st.selectbox("Select a topic", ["IT Act 2000", "Cybercrime Laws", "Digital Signature Regulations", "Data Protection", "E-commerce Regulations"])
    search_results = search_serper(f"Indian {cyber_topic} explanation")
    prompt = f"Provide a comprehensive explanation of {cyber_topic} in India based on the following search results with accurate and real time data. Provide the most important information in bullets and follow up with any other material that is relevant. Dictate the information as a news anchor and refrain from using terms like *here are the 10 bullets as asked*. Begin and end with actual relevant info.:\n\n{json.dumps(search_results['organic'][:5])}"
    explanation = query_groq(prompt)
    st.write(explanation)

def property_law_guide():
    st.subheader("Property Law Guide")
    property_topic = st.selectbox("Select a topic", ["Transfer of Property Act", "Real Estate Regulation Act (RERA)", "Rent Control Laws", "Land Acquisition Laws", "Stamp Duty and Registration"])
    search_results = search_serper(f"Indian {property_topic} explanation")
    prompt = f"Provide a comprehensive explanation of {property_topic} in India based on the following search results with accurate and real time data. Provide the most important information in bullets and follow up with any other material that is relevant. Dictate the information as a news anchor and refrain from using terms like *here are the 10 bullets as asked*. Begin and end with actual relevant info.:\n\n{json.dumps(search_results['organic'][:5])}"
    explanation = query_groq(prompt)
    st.write(explanation)

def legal_process_navigator():
    st.subheader("Legal Process Navigator")
    process = st.selectbox("Select a legal process", [
        "Filing a PIL", 
        "RTI Application", 
        "FIR Filing",
        "Filing a Consumer Complaint",
        "Applying for Bail",
        "Filing for Divorce",
        "Registering a Company",
        "Filing Income Tax Returns",
        "Applying for a Patent",
        "Filing a Writ Petition"
    ])
    search_results = search_serper(f"Steps for {process} in India")
    prompt = f"Provide a step-by-step guide for {process} in India based on the following search results:\n\n{json.dumps(search_results['organic'][:5])}"
    guide = query_groq(prompt)
    st.write(guide)

def document_templates():
    st.subheader("Legal Document Templates")
    template_type = st.selectbox("Select a template", [
        "Affidavit", 
        "Rental Agreement", 
        "Will", 
        "Power of Attorney", 
        "Sale Deed", 
        "Gift Deed", 
        "Employment Agreement", 
        "Non-Disclosure Agreement (NDA)", 
        "Divorce Petition", 
        "Partnership Deed"
    ])
    search_results = search_serper(f"{template_type} template India")
    prompt = f"Generate an official and legal template for a {template_type} in India based on the following search results with accurate and real time data:\n\n{json.dumps(search_results['organic'][:5])}"
    template = query_groq(prompt)
    st.write(template)

def legal_news():
    st.subheader("Legal News and Updates")
    search_results = search_serper("Latest Indian legal news")
    prompt = f"Summarize the following latest Indian legal news. Provide output as 10 bullet points that cover all the latest news at that time. Dictate the news as a news anchor and refrain from using terms like *here are the 10 bullets as asked*. Begin and end with actual relevant info. Make sure all news is related to indian law and legal system and as per the indian constitution:\n\n{json.dumps(search_results['organic'][:5])}"
    news_summary = query_groq(prompt)
    st.write(news_summary)

def legal_aid():
    st.subheader("Legal Aid Directory")
    location = st.text_input("Enter your location")
    specialization = st.selectbox("Select specialization", ["Criminal", "Civil", "Family", "Corporate"])
    if st.button("Search"):
        search_results = search_serper(f"Legal aid {specialization} lawyer in {location} India")
        prompt = f"Provide a list of legal aid options for {specialization} law in {location}, India based on the following search results:\n\n{json.dumps(search_results['organic'][:5])}"
        aid_results = query_groq(prompt)
        st.write(aid_results)

def case_tracker():
    st.subheader("Court Case Tracker")
    case_number = st.text_input("Enter case number")
    if st.button("Track Case"):
        search_results = search_serper(f"Indian court case status {case_number}")
        prompt = f"Provide the current status and summary of the Indian court case with number {case_number} based on the following search results:\n\n{json.dumps(search_results['organic'][:5])}"
        case_status = query_groq(prompt)
        st.write(case_status)

def create_docx(content):
    doc = Document()
    doc.add_heading('Legal Complaint', 0)
    
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        doc.add_paragraph(para)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def constitutional_history():
    st.subheader("Constitutional History of India")
    topic = st.text_input("Enter a topic related to Indian law and order history")
    if st.button("Generate History Lesson"):
        search_results = search_serper(f"Indian constitutional history: {topic}")
        prompt = f"Generate a concise history lesson about {topic} in the context of Indian law and order. Focus on key dates, events, significant figures, and lasting impacts. Pros/cons of the topic. Keep the response brief and informative and include key analysis in bullets. Dictate the information as a news anchor and refrain from using terms like *here are the 10 bullets as asked*. Begin and end with actual relevant info.:\n\n{json.dumps(search_results['organic'][:3])}"
        history_lesson = query_groq(prompt)
        st.write(history_lesson)


if __name__ == '__main__':
    main()